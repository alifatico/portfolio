#!/usr/bin/env python3
"""
Extract structured role data from a .docx handover document.
Outputs JSON: document title, section headings, list-item responsibilities, detected tools.
Usage: python3 scripts/extract_role.py <path-to-docx>
"""
import sys
import json
import re

try:
    from docx import Document
except ImportError:
    print("pip install python-docx", file=sys.stderr)
    sys.exit(1)

TOOL_KEYWORDS = {
    "Excel", "VBA", "Power BI", "PowerPoint", "SharePoint", "ThinkCell",
    "SAP", "Python", "SQL", "Microsoft Teams", "Outlook", "Tableau", "Salesforce",
    "Git", "Apify",
}


def extract(path: str) -> dict:
    doc = Document(path)
    paras = [p for p in doc.paragraphs if p.text.strip()]

    title = paras[0].text.strip() if paras else ""
    subtitle = paras[1].text.strip() if len(paras) > 1 else ""

    headings = [
        p.text.strip()
        for p in doc.paragraphs
        if "heading" in p.style.name.lower() and p.text.strip()
    ]

    responsibilities = [
        p.text.strip()
        for p in doc.paragraphs
        if "list" in p.style.name.lower() and p.text.strip()
    ]

    all_text = " ".join(p.text for p in doc.paragraphs)
    found_tools = sorted(
        {t for t in TOOL_KEYWORDS if re.search(rf"\b{re.escape(t)}\b", all_text, re.IGNORECASE)}
    )

    return {
        "source": str(path),
        "document_title": title,
        "document_subtitle": subtitle,
        "section_headings": headings,
        "responsibilities_raw": responsibilities[:25],
        "detected_tools": found_tools,
    }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/extract_role.py <path-to-docx>")
        sys.exit(1)
    result = extract(sys.argv[1])
    print(json.dumps(result, indent=2, ensure_ascii=False))
